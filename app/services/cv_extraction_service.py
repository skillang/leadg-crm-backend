# app/services/cv_extraction_service.py - Core CV Text Extraction Service

import io
import re
import fitz  # PyMuPDF
import docx
import dateutil.parser
from datetime import datetime
from typing import Dict, Any, Optional, Tuple, List
import logging
import time

logger = logging.getLogger(__name__)

class CVExtractionService:
    """Service for extracting structured data from CV files"""
    
    def __init__(self):
        self.supported_types = {
            'application/pdf': ['.pdf'],
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': ['.docx'],
            'application/msword': ['.doc']
        }
        self.max_file_size = 10 * 1024 * 1024  # 10MB
        self.extractor_version = "1.0"
        
        # Confidence thresholds
        self.high_confidence = 0.85
        self.medium_confidence = 0.60
        self.low_confidence = 0.40
    
    def validate_file(self, file_content: bytes, filename: str, mime_type: str) -> Dict[str, Any]:
        """Validate uploaded CV file"""
        validation_result = {
            "valid": True,
            "errors": [],
            "warnings": [],
            "file_info": {
                "size_bytes": len(file_content),
                "size_mb": round(len(file_content) / (1024 * 1024), 2),
                "filename": filename,
                "mime_type": mime_type
            }
        }
        
        # Check file size
        if len(file_content) > self.max_file_size:
            validation_result["valid"] = False
            validation_result["errors"].append(f"File size ({validation_result['file_info']['size_mb']}MB) exceeds 10MB limit")
        
        # Check file type
        if mime_type not in self.supported_types:
            validation_result["valid"] = False
            validation_result["errors"].append(f"File type '{mime_type}' not supported. Only PDF and DOCX files are allowed.")
        
        # Check if file has content
        if len(file_content) == 0:
            validation_result["valid"] = False
            validation_result["errors"].append("File is empty")
        
        # File size warnings
        if len(file_content) > 5 * 1024 * 1024:  # 5MB
            validation_result["warnings"].append("Large file size may result in slower processing")
        
        return validation_result
    
    def extract_text_from_file(self, file_content: bytes, mime_type: str, filename: str = "") -> Tuple[str, Dict[str, Any]]:
        """Extract raw text from CV file with metadata"""
        start_time = time.time()
        
        try:
            if mime_type == 'application/pdf':
                text, file_info = self._extract_text_from_pdf(file_content)
            elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
                text, file_info = self._extract_text_from_docx(file_content)
            else:
                raise Exception(f"Unsupported file type: {mime_type}")
            
            processing_time = (time.time() - start_time) * 1000  # Convert to milliseconds
            
            file_info.update({
                "processing_time_ms": round(processing_time, 2),
                "text_length": len(text),
                "extractor_version": self.extractor_version
            })
            
            return text, file_info
            
        except Exception as e:
            raise Exception(f"Error extracting text from {filename}: {str(e)}")
    
    def _extract_text_from_pdf(self, file_content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF file"""
        try:
            doc = fitz.open(stream=file_content, filetype="pdf")
            text_parts = []
            
            file_info = {
                "pages": len(doc),
                "file_type": "PDF"
            }
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                page_text = page.get_text()
                if page_text.strip():
                    text_parts.append(page_text)
            
            doc.close()
            text = "\n".join(text_parts)
            
            if not text.strip():
                raise Exception("No text content found in PDF")
            
            return text, file_info
            
        except Exception as e:
            raise Exception(f"Error processing PDF: {str(e)}")
    
    def _extract_text_from_docx(self, file_content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(io.BytesIO(file_content))
            paragraphs = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            
            text = "\n".join(paragraphs)
            
            file_info = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "file_type": "DOCX"
            }
            
            if not text.strip():
                raise Exception("No text content found in DOCX")
            
            return text, file_info
            
        except Exception as e:
            raise Exception(f"Error processing DOCX: {str(e)}")
    
    def extract_name(self, text: str) -> Tuple[str, float]:
        """Extract name with confidence score"""
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        if not lines:
            return "", 0.0
        
        # Skip patterns that are unlikely to be names
        skip_patterns = [
            r'resume|curriculum|vitae|cv|personal\s+information',
            r'email|phone|mobile|address|contact',
            r'objective|summary|profile|experience|education|skills',
            r'www\.|http|\.com|\.org|@',
            r'^\d+',
            r'[<>@#$%^&*(){}[\]]'
        ]
        
        # Name patterns with confidence scoring
        name_patterns = [
            (r'^([A-Z][a-z]+ [A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)$', 0.98),
            (r'(?:CURRICULUM VITAE|CV|RESUME)\s*\n+([A-Z][a-z]+ [A-Z][a-z]+)', 0.95),
            (r'^([A-Z][A-Z\s]+)$', 0.85),
            (r'^([A-Za-z]+\s+[A-Za-z]+(?:\s+[A-Za-z]+)?)$', 0.80),
        ]
        
        # Check first few lines for name
        for line in lines[:10]:
            # Skip if matches skip patterns
            if any(re.search(pattern, line, re.IGNORECASE) for pattern in skip_patterns):
                continue
            
            # Check against name patterns
            for pattern, confidence in name_patterns:
                match = re.search(pattern, line)
                if match:
                    name = self._clean_name(match.group(1))
                    if len(name.split()) >= 2:
                        return name, confidence
        
        # Additional extraction for specific resume patterns
        cv_pattern = r'(?:CURRICULUM VITAE|CV|RESUME)\s*\n+([A-Za-z\s]+?)(?:\n|E-mail|Email|Contact)'
        match = re.search(cv_pattern, text, re.IGNORECASE | re.MULTILINE)
        if match:
            name = self._clean_name(match.group(1))
            if len(name.split()) >= 2:
                return name, 0.90
        
        # Fallback: return first non-empty line that looks like a name
        for line in lines[:5]:
            if not any(re.search(pattern, line, re.IGNORECASE) for pattern in skip_patterns):
                cleaned = self._clean_name(line)
                if 2 <= len(cleaned.split()) <= 4:
                    return cleaned[:50], 0.5
        
        return "Name not found", 0.2
    
    def extract_email(self, text: str) -> Tuple[str, float]:
        """Extract email with confidence score"""
        email_patterns = [
            (r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', 0.95),
            (r'(?:Email|E-mail|E-mail id):\s*([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,})', 0.98),
            (r'(?:PERSONAL INFORMATION|Email)\s*\n*[^\n]*?([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,})', 0.92),
        ]
        
        exclude_domains = ['example.com', 'test.com', 'domain.com', 'email.com', 'mail.com']
        
        for pattern, confidence in email_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for email in matches:
                email = email.lower().strip()
                domain = email.split('@')[1] if '@' in email else ''
                
                if domain not in exclude_domains and '.' in domain and len(domain) > 3:
                    return email, confidence
        
        return "", 0.0
    
    def extract_phone(self, text: str) -> Tuple[str, float]:
        """Extract phone number with confidence score"""
        phone_patterns = [
            (r'\(\+91\)\s*(\d{10})', 0.95),
            (r'\+91[-.\s]?(\d{10})', 0.95),
            (r'(?:Contact No|Mobile|Phone):\s*\+?91[-.\s]?(\d{10})', 0.98),
            (r'\b([6-9]\d{9})\b', 0.80),
            (r'\+\d{1,3}[-.\s]?\(?\d{3,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{4}', 0.90),
            (r'\(\d{3,4}\)\s*\d{3,4}[-.\s]?\d{4}', 0.85),
        ]
        
        for pattern, confidence in phone_patterns:
            matches = re.findall(pattern, text)
            if matches:
                phone = matches[0] if isinstance(matches[0], str) else matches[0]
                phone = re.sub(r'[^\d+]', '', phone)
                if 10 <= len(phone.replace('+', '')) <= 15:
                    if len(phone) == 10 and phone[0] in ['6', '7', '8', '9']:
                        phone = '+91' + phone
                    return phone, confidence
        
        return "", 0.0
    
    def extract_age(self, text: str) -> Tuple[Optional[int], float]:
        """Extract age with confidence score"""
        age_patterns = [
            (r'age\s*[:\-]?\s*(\d{2})', 0.90),
            (r'(\d{2})\s*years?\s*old', 0.85),
            (r'age\s*(\d{2})', 0.80),
        ]
        
        for pattern, confidence in age_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                age = int(match)
                if 18 <= age <= 65:
                    return age, confidence
        
        # Date of birth patterns
        dob_patterns = [
            r'(?:dob|date\s*of\s*birth|born)\s*[:\-]?\s*(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})',
            r'(?:dob|born)\s*[:\-]?\s*(\d{1,2}\s+(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\w*\s+\d{2,4})',
        ]
        
        for pattern in dob_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                try:
                    parsed_date = dateutil.parser.parse(match)
                    if 1940 <= parsed_date.year <= 2010:
                        age = datetime.now().year - parsed_date.year
                        return age, 0.75
                except:
                    continue
        
        return None, 0.0
    
    def extract_skills(self, text: str) -> Tuple[str, float]:
        """Extract skills with confidence score"""
        skills_patterns = [
            (r'(?:KEY\s*SKILLS|SKILLS|COMPETENCIES|TECHNOLOGIES)[:\-]?\s*([^\n]*(?:\n[^\n]*){0,8})', 0.90),
            (r'(?:TECHNICAL\s*SKILLS|PROFESSIONAL\s*SKILLS)[:\-]?\s*([^\n]*(?:\n[^\n]*){0,5})', 0.85),
        ]
        
        for pattern, base_confidence in skills_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                skills_text = match.group(1)
                skills = re.split(r'[,;\n\t]', skills_text)
                skills = [s.strip() for s in skills if s.strip() and len(s.strip()) > 1]
                
                exclude_words = ['years', 'experience', 'knowledge', 'familiar', 'working', 'other', 'personal', 'details', 'including']
                skills = [s for s in skills if not any(e in s.lower() for e in exclude_words) and len(s) < 50]
                
                if skills:
                    return ', '.join(skills[:12]), base_confidence
        
        # Look for medical/nursing specific skills
        medical_keywords = re.findall(
            r'\b(?:ICU|Medical|Surgical|Nursing|BSc|BLS|ECMO|CRRT|Ventilator|Cardiac|Intensive Care|Emergency|Critical Care|Patient Care|NABH|JCI|Multispeciality|TL|Team Leader|Life Support|Defibrillator|Corona Warrior|Staff Nurse|Registered Nurse)\b', 
            text, 
            re.IGNORECASE
        )
        
        # Look for technical keywords
        tech_keywords = re.findall(
            r'\b(?:Python|Java|JavaScript|React|Node\.?js|Angular|Vue|PHP|C\+\+|C#|Ruby|Go|Swift|Kotlin|HTML|CSS|SQL|MongoDB|PostgreSQL|MySQL|AWS|Azure|Docker|Kubernetes|Git|Linux|Windows|Computer|Internet|Basic|Knowledge)\b', 
            text, 
            re.IGNORECASE
        )
        
        # Look for general professional skills
        professional_keywords = re.findall(
            r'\b(?:Management|Leadership|Team|Communication|Problem Solving|Critical Thinking|Training|Supervision|Patient Care|Quality|Safety|Compliance|Documentation|Coordination)\b',
            text, 
            re.IGNORECASE
        )
        
        all_keywords = medical_keywords + tech_keywords + professional_keywords
        if all_keywords:
            unique_skills = list(dict.fromkeys([skill.title() for skill in all_keywords]))
            return ', '.join(unique_skills[:12]), 0.75
        
        return "", 0.0
    
    def extract_education(self, text: str) -> Tuple[str, float]:
        """Extract education with confidence score"""
        education_patterns = [
            (r'(?:education|academic|qualification)[:\-]?\s*([^\n]*(?:\n[^\n]*){0,5})', 0.85),
            (r'((?:bachelor|master|phd|doctorate|diploma|certificate|degree|b\.?sc|m\.?sc|b\.?a|m\.?a|b\.?tech|m\.?tech|mba|bca|mca|be|me|ms|bs)\s+[^\n]*)', 0.90),
            (r'(university|college|institute|school)\s*[:\-]?\s*([^\n]*(?:\n[^\n]*){0,3})', 0.75),
        ]
        
        best_education = ""
        best_confidence = 0.0
        
        for pattern, confidence in education_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
            for match in matches:
                if isinstance(match, tuple):
                    edu_text = ' '.join(match).strip()
                else:
                    edu_text = match.strip()
                
                edu_text = ' '.join(edu_text.split()[:25])
                edu_text = re.sub(r'\d{4}', '', edu_text)
                edu_text = edu_text.strip()
                
                if len(edu_text) > 10 and confidence > best_confidence:
                    best_education = edu_text
                    best_confidence = confidence
        
        return best_education, best_confidence
    
    def extract_experience(self, text: str) -> Tuple[str, float]:
        """Extract experience with confidence score"""
        experience_patterns = [
            (r'(?:Total work experience|Work experience|Experience):\s*([^.\n]+)', 0.95),
            (r'(\d+\+?\s*[Yy]ears?\s*\d*\s*[Mm]onths?)', 0.90),
            (r'(\d+\+?\s*years?\s+experience)', 0.85),
            (r'(?:worked|working)\s*(?:as|at|for)\s*([^\n]+)', 0.70),
            (r'(?:total\s*work\s*experience|work\s*experience|experience)[:\-]?\s*([^\n]*(?:\n[^\n]*){0,2})', 0.80),
        ]
        
        best_experience = ""
        best_confidence = 0.0
        
        for pattern, confidence in experience_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
            for match in matches:
                exp_text = match.strip()
                exp_text = ' '.join(exp_text.split()[:15])
                
                if len(exp_text) > 3 and confidence > best_confidence:
                    best_experience = exp_text
                    best_confidence = confidence
        
        # Look in profile summary for experience mentions
        profile_match = re.search(r'(?:PROFILE SUMMARY|SUMMARY|OBJECTIVE)(.*?)(?:EDUCATION|WORK EXPERIENCE|KEY SKILLS)', 
                                 text, re.IGNORECASE | re.DOTALL)
        if profile_match:
            profile_text = profile_match.group(1)
            exp_in_profile = re.search(r'(\d+\+?\s*years?\s+experience[^.]*)', profile_text, re.IGNORECASE)
            if exp_in_profile and best_confidence < 0.88:
                best_experience = exp_in_profile.group(1)
                best_confidence = 0.88
        
        return best_experience, best_confidence
    
    def extract_all_details(self, text: str, filename: str) -> Dict[str, Any]:
        """Extract all structured data from CV text"""
        extraction_start = time.time()
        
        # Extract all fields with confidence scores
        name, name_confidence = self.extract_name(text)
        email, email_confidence = self.extract_email(text)
        phone, phone_confidence = self.extract_phone(text)
        age, age_confidence = self.extract_age(text)
        skills, skills_confidence = self.extract_skills(text)
        education, education_confidence = self.extract_education(text)
        experience, experience_confidence = self.extract_experience(text)
        
        # Build extraction result
        extracted_data = {
            'name': name,
            'email': email,
            'phone': phone,
            'age': age,
            'skills': skills,
            'education': education,
            'experience': experience,
        }
        
        # Build confidence scores
        confidence_scores = {
            'name': name_confidence,
            'email': email_confidence,
            'phone': phone_confidence,
            'age': age_confidence,
            'skills': skills_confidence,
            'education': education_confidence,
            'experience': experience_confidence,
        }
        
        # Calculate overall metrics
        processing_time = (time.time() - extraction_start) * 1000
        
        # Calculate overall confidence (weighted average)
        weights = {'name': 0.20, 'email': 0.25, 'phone': 0.15, 'skills': 0.20, 'education': 0.10, 'experience': 0.10}
        overall_confidence = sum(confidence_scores[field] * weights[field] for field in weights.keys())
        
        # Calculate field completeness
        filled_fields = sum(1 for value in extracted_data.values() if value is not None and str(value).strip())
        field_completeness = filled_fields / len(extracted_data)
        
        # Identify quality issues
        quality_issues = []
        for field, confidence in confidence_scores.items():
            if confidence < self.low_confidence:
                quality_issues.append(f"Low confidence extraction for {field}")
            elif confidence < self.medium_confidence:
                quality_issues.append(f"Medium confidence extraction for {field}")
        
        if not email:
            quality_issues.append("No email address found")
        if not phone:
            quality_issues.append("No phone number found")
        
        return {
            'extracted_data': extracted_data,
            'confidence_scores': confidence_scores,
            'extraction_metadata': {
                'original_filename': filename,
                'processing_time_ms': round(processing_time, 2),
                'raw_text_length': len(text),
                'extractor_version': self.extractor_version,
                'extraction_timestamp': datetime.utcnow(),
                'overall_confidence': round(overall_confidence, 3),
                'field_completeness': round(field_completeness, 3),
                'quality_issues': quality_issues,
                'recommended_review': overall_confidence < self.medium_confidence or len(quality_issues) > 2
            }
        }
    
    def _clean_name(self, name: str) -> str:
        """Clean and format extracted name"""
        if not name:
            return ""
        
        # Remove special characters and extra spaces
        name = re.sub(r'[^\w\s\.]', '', name).strip()
        name = ' '.join(name.split())  # Remove extra spaces
        
        # Filter out common non-name words
        exclude_words = ['resume', 'cv', 'curriculum', 'vitae', 'contact', 'phone', 'email', 'mobile', 'total', 'work', 'page', 'document']
        words = [word for word in name.split() if word.lower() not in exclude_words]
        
        cleaned_name = ' '.join(words) if words else name
        
        # Capitalize properly
        return ' '.join(word.capitalize() for word in cleaned_name.split())
    
    def calculate_extraction_quality_score(self, confidence_scores: Dict[str, float], extracted_data: Dict[str, Any]) -> float:
        """Calculate overall extraction quality score (0-100)"""
        # Base score from confidence levels
        weighted_confidence = (
            confidence_scores.get('name', 0) * 0.20 +
            confidence_scores.get('email', 0) * 0.25 +
            confidence_scores.get('phone', 0) * 0.15 +
            confidence_scores.get('skills', 0) * 0.20 +
            confidence_scores.get('education', 0) * 0.10 +
            confidence_scores.get('experience', 0) * 0.10
        )
        
        base_score = weighted_confidence * 70  # Max 70 from confidence
        
        # Bonus for data completeness
        completeness_bonus = 0
        required_fields = ['name', 'email', 'phone']
        optional_fields = ['skills', 'education', 'experience', 'age']
        
        for field in required_fields:
            if extracted_data.get(field):
                completeness_bonus += 10  # 10 points per required field
        
        for field in optional_fields:
            if extracted_data.get(field):
                completeness_bonus += 2.5  # 2.5 points per optional field
        
        total_score = min(base_score + completeness_bonus, 100)
        return round(total_score, 1)
    
    def get_extraction_recommendations(self, confidence_scores: Dict[str, float], extracted_data: Dict[str, Any]) -> List[str]:
        """Generate recommendations for improving extraction quality"""
        recommendations = []
        
        # Check for missing critical fields
        if not extracted_data.get('email'):
            recommendations.append("Consider manually adding email address")
        if not extracted_data.get('phone'):
            recommendations.append("Consider manually adding phone number")
        
        # Check confidence levels
        low_confidence_fields = [field for field, confidence in confidence_scores.items() if confidence < self.medium_confidence]
        if low_confidence_fields:
            recommendations.append(f"Review and verify: {', '.join(low_confidence_fields)}")
        
        # Check data quality
        if extracted_data.get('skills') and len(extracted_data['skills'].split(',')) < 3:
            recommendations.append("Consider adding more specific skills")
        
        if not extracted_data.get('experience'):
            recommendations.append("Consider adding work experience details")
        
        return recommendations